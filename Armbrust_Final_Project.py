# Greta Armbrust
# Final Project
# Comp112-01

"""
1. Flow Control, File IO, Working with external library

2. I believe that through the completion of this code I have shown that I have a good grasp of basic concepts like loops which we learned early on as well as the ability to add to that knowledge through the learning of an external dictionary.
    While I believe that mastery may be being generous I certainly have a good grasp on it and would feel comfortable creating more things in python particularly those which can be aligned with my other interests.

3. In this class I have learned a good amount of introductory python as well as developing an ability to expand on the knowledge should I wish
"""

import os

import docx

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text import font

error_message = 'Error. Not a valid submission. Please run program again.'
thanks = 'Thank you.'

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Courier'
font.size = Pt(12)

def line_by_line(label, text, document):
    """
    sig: str (three of them) -> NoneType

    Takes the labelled text and prints it into the doc properly formatted
    """
                    
    if label=='HEAD':
                            
        text = text.upper()
                            
        paragraph = document.add_paragraph(text)
        paragraph.style = document.styles['Normal']

        paragraph_format = paragraph.paragraph_format

    elif label=='ACTION':

        paragraph = document.add_paragraph(text)
        paragraph.style = document.styles['Normal']

    elif label=='CHAR':

        text = text.upper()

        paragraph = document.add_paragraph(text)
        paragraph.style = document.styles['Normal']

        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(2.0)
                            
    elif label=='DIA':
                            
        paragraph = document.add_paragraph(text)
        paragraph.style = document.styles['Normal']
                            
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(1.0)
        paragraph_format.right_indent = Inches(1.5)

    elif label=='PAREN':

        paragraph = document.add_paragraph('(' + text + ')')
        paragraph.style = document.styles['Normal']
                            
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(1.5)
        paragraph_format.right_indent = Inches(2.0)

    elif label=='TRAN':

        text = text.upper()

        paragraph = document.add_paragraph(text +':')
        paragraph.style = document.styles['Normal']
                            
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(4.0)

    elif label=='SHOT':

        text = text.upper()

        paragraph = document.add_paragraph(text +'-')
        paragraph.style = document.styles['Normal']

def screenplay_format():
    """
    sig: NoneType -> NoneType
    sets up the text and document and either creates a new one of adds to a pre-existing one.
    """

    print('We will now format your text.')
    start_text = input('Please input the name of this text. ')
    print(thanks)
                      
    ans = input('Should the screenplay be saved as a new document of appended to an old one? ')
    print()
                
    if ans == 'new' or ans == 'New':
        title = input('What is the title of the screenplay? ')
        file_name = title + '.docx'

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.0)

    elif ans == 'append' or ans == 'Append':
        print('Please make sure the file you wish to append is whithin this folder')

        file_name = input ("Please input the pre-existing screenplay's filename")

    else:
        print (error_message)

    start_text = open(start_text, 'r')
    lines = start_text.readlines()
    for line in lines:
        label, text = line.split(':', 1)
        text = text.strip()
            
        line_by_line(label, text, document)
        
    document.save(file_name)
            
def screenplay_final_project():
    """
    sig: NoneType -> NoneType
    this function creates the structure of the code and contains the full first and third segment which allows the user to look up information and to create a title page 
    """

    print ('Hello, Thank you for using this program.')
    print ('Please select what you would like me to do.')
    print ('A. Ask for formatting for specific screenplay elements')
    print ('B. Format prewritten text')
    print ('C. Make a title page')
    choice = input ('Type the letter of your choice here: ')
    print()

    if choice == 'a' or choice == 'A' : 
        element = input ('What element would you like formatting info for? ')
        print()
        
        all_caps = 'It will be in all caps.'
        sentence_case = 'Printed like standard text.'
        no_indents = 'There are no indents used.'
        shooting_script = ('This generally only appears in shooting scripts')
        
        if element == 'Scene Heading':
            print(no_indents)
            print(all_caps)
            print('Example:')
            print('EXT. LIBRARY - DAY')
            print('This means a scene is taking place outside of a library during the day.')
        elif element == 'Action':
            print('Description of a scene which is written in present tense.')
            print(no_indents)
            print(sentence_case)
        elif element == 'Character':
            print('Name of the character placed above dialogue')
            print('Indent 2" on the left')
            print(all_caps)
        elif element == 'Dialogue':
            print('Indent 1" on the left and 1.5" on the right')
            print(sentence_case)
        elif element == 'Parenthetical':
            print('Placed between charater name and dialogue')
            print('Indent 1.5" on the left and 2" on the right')
            print (sentence_case)
        elif element == 'Transition':
            print('These are editing directions')
            print(shooting_script)
            print('Indent 4" on the left')
            print(all_caps)
        elif element == 'Shot':
            print('These specify particular shots')
            print(shooting_script)
            print(no_indents)
            print(all_caps)
        else:
            print (error_message)

    elif choice == 'b' or choice == 'B' :
        print('Before we run this please double check that your text is properly coded.')
        ans = input('Would you like to review the formatting codes? ')
        print()
        
        if ans == 'yes' or ans == 'Yes' or ans == 'Y' or ans == 'y':
            
            folder_location = input('Please input the location of this program on your computer. ')

            print()
            
            format_codes = open(folder_location + '/Formatting_Codes.txt','r')
            for line in format_codes:
                print(line)
            format_codes.close
            
            screenplay_format()
            
        elif ans == 'no' or ans == 'No' or ans == 'N' or ans == 'n':
            
            screenplay_format()
            
        else:
            print (error_message)

    elif choice == 'c' or choice == 'C':
        print ('We will be formating a title page')

        title = input('What is the title of the screenplay this page is for? ')

        file_name = title + ' Title Page.docx'
        
        document.save(file_name)

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(4.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.0)

        title = title.upper()

        paragraph = document.add_paragraph(title)
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = document.add_paragraph()


        paragraph = document.add_paragraph('written by')
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name = input ('What is your name? ')

        paragraph = document.add_paragraph(name)
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = document.add_paragraph('\r')
        paragraph = document.add_paragraph('\r')
        paragraph = document.add_paragraph('\r')
        paragraph = document.add_paragraph('\r')
        paragraph = document.add_paragraph('\r')
        paragraph = document.add_paragraph('\r')

        email = input ('What is your email? ')

        paragraph = document.add_paragraph(email)
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        phone = input ('What is your phone number? ')
        
        paragraph = document.add_paragraph(phone)
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        contact = input ('Please put any additional lines of contact information here. ')

        paragraph = document.add_paragraph(contact)
        paragraph.style = document.styles['Normal']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        document.save(file_name)
        
    else:
        print (error_message)
        
screenplay_final_project()
